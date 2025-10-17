import os
import shutil
import json
import requests
from pathlib import Path
from datetime import datetime
import pandas as pd
from docx import Document
from pdf2docx import Converter
from docxcompose.composer import Composer
from pymongo import MongoClient
from docx.table import Table
from docx.text.paragraph import Paragraph
from concurrent.futures import ThreadPoolExecutor, as_completed
from aiTaskRunner import AITaskRunner
from typing import Optional

RESULT_DIR = Path(r'files\result')
RESULT_DIR.mkdir(parents=True, exist_ok=True)


# ======================== 工具函数 ========================
def make_unique_subdir(base_dir: str, prefix: str = "folder") -> Path:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    subdir = Path(base_dir) / f"{prefix}_{timestamp}"
    subdir.mkdir(parents=True, exist_ok=True)
    return subdir


def make_unique_filename(dir_path: str, filename: str) -> Path:
    dir_path = Path(dir_path)
    dir_path.mkdir(parents=True, exist_ok=True)

    file_path = dir_path / filename
    counter = 1
    while file_path.exists():
        file_path = dir_path / f"{Path(filename).stem}_{counter}{Path(filename).suffix}"
        counter += 1
    return file_path


# ======================== 文件下载 ========================
def download_files_to_unique_folder(urls: list,
                                    base_dir=r"files\01-提取的文件",
                                    max_workers=5) -> str:
    input_dir = make_unique_subdir(base_dir, prefix="download")
    SUPPORTED_TYPES = {
        'application/pdf': '.pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
        'application/msword': '.doc',
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
        'application/vnd.ms-excel': '.xls',
        'text/csv': '.csv'
    }

    downloaded_files = []

    def download_single(url, idx):
        try:
            print(f"[下载] {idx+1}: {url}")
            r = requests.get(url, stream=True, timeout=30)
            r.raise_for_status()
            content_type = r.headers.get('content-type', '').split(';')[0]
            ext = SUPPORTED_TYPES.get(content_type) or Path(url).suffix.lower()
            if not ext or ext not in SUPPORTED_TYPES.values():
                print(f"[跳过] 不支持文件类型: {url}")
                return None

            filename = Path(url).name or f"file_{idx+1}{ext}"
            file_path = make_unique_filename(input_dir, filename)

            with open(file_path, 'wb') as f:
                for chunk in r.iter_content(65536):
                    if chunk:
                        f.write(chunk)
            print(f"[成功] {file_path}")
            return str(file_path)
        except Exception as e:
            print(f"[下载失败] {url} - {e}")
            return None

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_url = {executor.submit(download_single, url, i): url for i, url in enumerate(urls)}
        for future in as_completed(future_to_url):
            result = future.result()
            if result:
                downloaded_files.append(result)

    print(f"[完成] 共下载 {len(downloaded_files)} 个文件，保存到 {input_dir}")
    return str(input_dir)


# ======================== 文件转换 Word ========================
def convert_files_to_docx_unique_folder(input_dir,
                                        base_output_dir=r"files\02-格式转换后的文件") -> str:
    output_dir = make_unique_subdir(base_output_dir, prefix="converted")
    input_dir = Path(input_dir)
    output_files = []

    for file_path in input_dir.iterdir():
        if not file_path.is_file():
            continue

        ext = file_path.suffix.lower()
        filename = file_path.stem
        out_file = make_unique_filename(output_dir, f"{filename}.docx")

        try:
            if ext in [".xls", ".xlsx"]:
                xls = pd.ExcelFile(file_path)
                doc = Document()

                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    df = df.dropna(axis=1, how="all").dropna(axis=0, how="all").fillna("")
                    if df.empty:
                        continue

                    first_cell = str(df.iloc[0, 0])
                    if "附件" in first_cell or "标准表" in first_cell:
                        doc.add_heading(first_cell, level=1)
                        df = df.iloc[1:]
                    if df.empty:
                        continue

                    table = doc.add_table(rows=1, cols=len(df.columns))
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(df.columns):
                        hdr_cells[i].text = str(col)
                        if hdr_cells[i].paragraphs[0].runs:
                            hdr_cells[i].paragraphs[0].runs[0].bold = True

                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val)

                    doc.add_paragraph("")
                doc.save(out_file)

            elif ext == ".csv":
                df = pd.read_csv(file_path)
                df = df.dropna(axis=1, how="all").dropna(axis=0, how="all").fillna("")
                if df.empty:
                    continue

                doc = Document()
                table = doc.add_table(rows=1, cols=len(df.columns))
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(df.columns):
                    hdr_cells[i].text = str(col)
                    if hdr_cells[i].paragraphs[0].runs:
                        hdr_cells[i].paragraphs[0].runs[0].bold = True

                for _, row in df.iterrows():
                    row_cells = table.add_row().cells
                    for i, val in enumerate(row):
                        row_cells[i].text = str(val)

                doc.save(out_file)

            elif ext == ".pdf":
                cv = Converter(str(file_path))
                cv.convert(str(out_file), start=0, end=None)
                cv.close()

            elif ext in [".doc", ".docx"]:
                shutil.copy(str(file_path), str(out_file))

            else:
                print(f"[跳过] 不支持文件类型: {file_path}")
                continue

            output_files.append(str(out_file))
            print(f"[已处理] {file_path} -> {out_file}")

        except Exception as e:
            print(f"[转换失败] {file_path} - {e}")

    print(f"[完成] 共转换 {len(output_files)} 个文件，保存到 {output_dir}")
    return str(output_dir)


# ======================== Word 合并 ========================
def merge_word_documents(word_files, base_output_dir=r"files\03-合并后的word文件") -> str:
    if not word_files:
        print("[提示] 无 Word 文件可合并")
        return ""

    output_dir = make_unique_subdir(base_output_dir, prefix="merged")
    out_file = make_unique_filename(output_dir, "merged.docx")

    master = Document(word_files[0])
    composer = Composer(master)
    for f in word_files[1:]:
        composer.append(Document(f))
    composer.save(str(out_file))

    print(f"[完成] 合并文件: {out_file}")
    return str(out_file)

def extract_title_with_ai(docx_file: str, api_key: str,
                          prompt_file: str = r"prompt\提取标题.txt",
                          base_url: str = "https://api.deepseek.com",
                          model: str = "deepseek-chat",
                          batch_id: Optional[int] = None) -> dict:
    try:
        doc = Document(docx_file)
        all_text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
        if not all_text:
            return {"PaperName": [], "DataId": batch_id or "xxxxxx", "error": "文档为空或无文本"}
    except Exception as e:
        return {"PaperName": [], "DataId": batch_id or "xxxxxx", "error": f"无法读取 Word 文件: {e}"}

    try:
        with open(prompt_file, "r", encoding="utf-8") as f:
            template = f.read()
    except Exception as e:
        return {"PaperName": [], "DataId": batch_id or "xxxxxx", "error": f"无法读取提示词: {e}"}

    prompt = template.replace("{doc_text}", all_text)

    try:
        runner = AITaskRunner(api_key=api_key, base_url=base_url, model=model)
        raw_response = runner.run_task(prompt)
        print(f"[AI 原始响应]:\n{raw_response}")

        import re
        if isinstance(raw_response, str):
            cleaned = re.sub(r"```(?:json)?\n(.*?)```", r"\1", raw_response, flags=re.S).strip()
            result = json.loads(cleaned)
        else:
            result = raw_response

        # 替换 DataId 为实际 batch_id
        if batch_id is not None:
            result["DataId"] = batch_id

        safe_title = Path(docx_file).stem

        # === 使用唯一子文件夹保存 ===
        output_dir = make_unique_subdir(RESULT_DIR, prefix="titles")
        result_file = make_unique_filename(output_dir, f"{safe_title}_headings_ai.json")

        with open(result_file, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)

        print(f"[完成] AI 提取标题结果已保存: {result_file}")
        return result

    except Exception as e:
        return {"PaperName": [], "DataId": batch_id or "xxxxxx", "error": str(e)}
# ======================== Word 分块 ========================
class DocxSliceOptimized:
    def __init__(self, path: str):
        self.doc_path = path
        try:
            self.doc = Document(path)
        except Exception as e:
            raise ValueError(f"无法读取 Word 文件: {path}, 错误: {e}")

        self.blocks = []
        for block in self.doc.element.body:
            try:
                if block.tag.endswith('p'):
                    self.blocks.append(Paragraph(block, self.doc))
                elif block.tag.endswith('tbl'):
                    self.blocks.append(Table(block, self.doc))
            except Exception:
                continue

    def _block_to_text(self, block):
        try:
            if isinstance(block, Paragraph):
                return block.text.strip()
            elif isinstance(block, Table):
                return "\n".join(
                    "\t".join(cell.text.strip() for cell in row.cells)
                    for row in block.rows
                )
        except Exception:
            return ""
        return ""

    def chunk_json(self, chunk_size: int = 10,
                   output_dir: str = r"files/04-分块后的json",
                   output_file_name: str = None) -> str:
        out_dir = make_unique_subdir(output_dir, prefix="chunks")
        if not output_file_name:
            output_file_name = "chunks.json"
        json_path = make_unique_filename(out_dir, output_file_name)

        chunks = []
        index = 1
        for i in range(0, len(self.blocks), chunk_size):
            slice_blocks = self.blocks[i:i + chunk_size]
            content_lines = [self._block_to_text(b) for b in slice_blocks]
            content_lines = [line for line in content_lines if line.strip()]
            if not content_lines:
                continue
            content = "\n".join(content_lines)
            chunks.append({"index": index, "content": content})
            index += 1

        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(chunks, f, ensure_ascii=False, indent=2)

        print(f"[完成] 文档分块已保存至: {json_path}, 共 {len(chunks)} 块")
        return str(json_path)


def save_chunks_to_mongo_batch(chunks_json_path: str,
                               mongo_uri: str = "mongodb://localhost:27017",
                               db_name: str = "doc_db",
                               collection_name: str = "doc_batches"):
    """
    将 JSON 文件中的 chunks 保存为 MongoDB 新批次
    """
    try:
        with open(chunks_json_path, 'r', encoding='utf-8') as f:
            chunks = json.load(f)
        if not chunks:
            print("[提示] 无分块数据可保存")
            return None

        client = MongoClient(mongo_uri, serverSelectionTimeoutMS=5000)
        client.admin.command("ping")
        collection = client[db_name][collection_name]

        last_batch = collection.find_one(sort=[("batch_id", -1)])
        batch_id = (last_batch["batch_id"] + 1) if last_batch else 1

        record = {
            "batch_id": batch_id,
            "chunks": chunks,
            "time": datetime.now().isoformat()
        }
        collection.insert_one(record)
        client.close()

        print(f"[完成] 批次 {batch_id} 已保存，包含 {len(chunks)} 个分块")
        return batch_id

    except Exception as e:
        print(f"[错误] 保存到 MongoDB 失败: {e}")
        return None

# ======================== 流程入口 ========================
def process_pipeline_independent(urls: list, api_key: str):
    input_dir = download_files_to_unique_folder(urls)
    print("下载目录:", input_dir)

    converted_dir = convert_files_to_docx_unique_folder(input_dir)
    print("转换目录:", converted_dir)

    word_files = list(Path(converted_dir).glob("*.docx"))
    merged_file = merge_word_documents([str(f) for f in word_files])
    print("合并文件:", merged_file)

    chunks_json_path = None
    batch_id = None
    if merged_file:
        try:
            slicer = DocxSliceOptimized(merged_file)
            chunks_json_path = slicer.chunk_json()
            print("[完成] 文档分块 JSON 文件:", chunks_json_path)
        except Exception as e:
            print(f"[错误] 文档分块失败: {e}")

        try:
            if chunks_json_path:
                batch_id = save_chunks_to_mongo_batch(chunks_json_path)
        except Exception as e:
            print(f"[错误] 保存到 MongoDB 失败: {e}")

    titles_result = None
    if merged_file:
        try:
            titles_result = extract_title_with_ai(
                docx_file=merged_file,
                api_key=api_key,
                batch_id=batch_id  # 传入 batch_id
            )
            print("[完成] 标题提取结果:", titles_result)
        except Exception as e:
            print(f"[错误] 提取标题失败: {e}")

    return titles_result, batch_id


if __name__ == "__main__":
    urls = [
        "http://localhost:8000/files10.docx",
        "http://localhost:8000/files2.docx"
    ]
    api_key = ""

    titles_result, batch_id = process_pipeline_independent(urls, api_key)
    print("标题提取结果:", titles_result)
